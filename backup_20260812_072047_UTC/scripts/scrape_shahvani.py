#!/usr/bin/env python3
"""
Scrape shahvani.com/dastans, extract story links, fetch each story content,
and email as a Word document.
"""

import os
import sys
import requests
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from email.message import EmailMessage
import smtplib
import ssl
from datetime import datetime
from urllib.parse import urljoin
import time

BASE_URL = "https://shahvani.com"
LIST_URL = f"{BASE_URL}/dastans"
EMAIL_TO = "jozvahedi@gmail.com"
EMAIL_FROM = "jozvahedimohammadreza@gmail.com"  # Sender email

# SMTP config - using Gmail SMTP (needs app password)
SMTP_HOST = "smtp.gmail.com"
SMTP_PORT = 587
SMTP_USER = "jozvahedimohammadreza@gmail.com"
SMTP_PASS = os.environ.get("SMTP_PASS", "")

HEADERS = {
    "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36",
    "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,image/webp,*/*;q=0.8",
    "Accept-Language": "fa-IR,fa;q=0.9,en-US;q=0.8,en;q=0.7",
}


def fetch_page(url, session=None):
    """Fetch a page and return BeautifulSoup object."""
    sess = session or requests.Session()
    resp = sess.get(url, headers=HEADERS, timeout=30)
    resp.raise_for_status()
    return BeautifulSoup(resp.text, "html.parser")


def extract_story_links(soup):
    """Extract all /dastan/ links from the table in panel-body."""
    links = []
    # Find the table with story links
    for td in soup.select("td a[href^='/dastan/']"):
        href = td.get("href")
        title = td.get_text(strip=True)
        full_url = urljoin(BASE_URL, href)
        links.append({"url": full_url, "title": title})
    return links


def extract_story_content(soup):
    """Extract the main story content from .panel-body."""
    panel_body = soup.select_one("div.panel-body")
    if not panel_body:
        return None
    
    # Get text content, preserving some structure
    # Remove script/style tags
    for tag in panel_body(["script", "style", "nav", "header", "footer", "aside"]):
        tag.decompose()
    
    return panel_body.get_text(separator="\n", strip=True)


def create_word_document(stories, date_str):
    """Create a Word document with all stories."""
    doc = Document()
    
    # Title
    title = doc.add_heading(f"داستان‌های شهوانی - {date_str}", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Meta info
    p = doc.add_paragraph()
    p.add_run(f"تاریخ استخراج: {date_str}\n").italic = True
    p.add_run(f"تعداد داستان‌ها: {len(stories)}\n").italic = True
    p.add_run(f"منبع: {LIST_URL}\n").italic = True
    
    doc.add_paragraph("─" * 50)
    
    for i, story in enumerate(stories, 1):
        # Story title as heading
        heading = doc.add_heading(f"{i}. {story['title']}", level=1)
        
        # URL
        p = doc.add_paragraph()
        run = p.add_run(f"لینک: {story['url']}")
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0, 0, 255)
        run.italic = True
        
        # Content
        if story.get("content"):
            content_para = doc.add_paragraph(story["content"])
            content_para.paragraph_format.space_after = Pt(12)
        else:
            p = doc.add_paragraph("[محتوا قابل استخراج نبود]")
            p.runs[0].italic = True
        
        # Separator
        if i < len(stories):
            doc.add_paragraph("─" * 50)
    
    # Save to bytes
    import io
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def send_email(subject, body, attachment_bytes, filename):
    """Send email with Word attachment via Gmail SMTP."""
    if not SMTP_PASS:
        print("❌ SMTP_PASS environment variable not set!")
        print("   Please set it: export SMTP_PASS='your-app-password'")
        return False
    
    msg = EmailMessage()
    msg["Subject"] = subject
    msg["From"] = EMAIL_FROM
    msg["To"] = EMAIL_TO
    msg.set_content(body)
    
    # Add attachment
    msg.add_attachment(
        attachment_bytes.getvalue(),
        maintype="application",
        subtype="vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename=filename
    )
    
    context = ssl.create_default_context()
    try:
        # Use STARTTLS for port 587
        with smtplib.SMTP(SMTP_HOST, SMTP_PORT) as server:
            server.starttls(context=context)
            server.login(SMTP_USER, SMTP_PASS)
            server.send_message(msg)
        print(f"✅ Email sent to {EMAIL_TO}")
        return True
    except Exception as e:
        print(f"❌ Failed to send email: {e}")
        return False


def main():
    print(f"🕐 Starting scrape at {datetime.now().isoformat()}")
    print(f"📡 Fetching list page: {LIST_URL}")
    
    session = requests.Session()
    
    try:
        # Fetch list page
        list_soup = fetch_page(LIST_URL, session)
        
        # Extract story links
        story_links = extract_story_links(list_soup)
        print(f"📚 Found {len(story_links)} story links")
        
        if not story_links:
            print("⚠️ No stories found! Page structure may have changed.")
            return 1
        
        # Fetch each story
        stories = []
        for i, link in enumerate(story_links, 1):
            print(f"  [{i}/{len(story_links)}] Fetching: {link['title'][:50]}...")
            try:
                story_soup = fetch_page(link["url"], session)
                content = extract_story_content(story_soup)
                stories.append({
                    "title": link["title"],
                    "url": link["url"],
                    "content": content
                })
                time.sleep(0.5)  # Be polite
            except Exception as e:
                print(f"    ❌ Error: {e}")
                stories.append({
                    "title": link["title"],
                    "url": link["url"],
                    "content": f"[خطا در دریافت: {e}]"
                })
        
        # Create Word document
        date_str = datetime.now().strftime("%Y-%m-%d %H:%M")
        filename = f"shahvani_stories_{datetime.now().strftime('%Y%m%d')}.docx"
        
        print(f"📝 Creating Word document: {filename}")
        doc_buffer = create_word_document(stories, date_str)
        
        # Send email
        subject = f"داستان‌های شهوانی - {datetime.now().strftime('%Y-%m-%d')}"
        body = f"""سلام،

فایل ضمیمه حاوی {len(stories)} داستان از سایت شهوانی (بخش داستان‌ها) است.
تاریخ استخراج: {date_str}
منبع: {LIST_URL}

این ایمیل به صورت خودکار توسط کرون‌جاب هر روز ساعت ۱۰ صبح ارسال شده است.

با تشکر،
Hermes Agent
"""
        
        print(f"📧 Sending email to {EMAIL_TO}...")
        success = send_email(subject, body, doc_buffer, filename)
        
        if success:
            print("✅ Done!")
            return 0
        else:
            return 1
            
    except Exception as e:
        print(f"❌ Fatal error: {e}")
        import traceback
        traceback.print_exc()
        return 1


if __name__ == "__main__":
    sys.exit(main())