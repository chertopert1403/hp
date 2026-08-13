import os
import sys
import requests
from bs4 import BeautifulSoup
from docx import Document
from datetime import datetime
from urllib.parse import urljoin
import time
import subprocess

BASE_URL = "https://shahvani.com"
LIST_URL = f"{BASE_URL}/dastans"
REPO_DIR = "/tmp/hermes-backup-git"
GIT_REPO_URL = "https://chertopert1981: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@github.com/chertopert1981/hermes-backup.git"

# 1. Fetch stories
sess = requests.Session()
resp = sess.get(LIST_URL, headers={"User-Agent": "Mozilla/5.0"}, timeout=30)
soup = BeautifulSoup(resp.text, "html.parser")
stories = []
for td in soup.select("td a[href^='/dastan/']")[:10]: # Test first 10
    url = urljoin(BASE_URL, td.get("href"))
    stories.append({"title": td.get_text(strip=True), "url": url})

# 2. Create Word
doc = Document()
doc.add_heading("Shahvani Stories Test", 0)
for s in stories:
    doc.add_heading(s['title'], level=1)
    doc.add_paragraph(s['url'])
filename = "shahvani_stories_test.docx"
doc.save(filename)

# 3. Git Push
if not os.path.exists(REPO_DIR):
    subprocess.run(["git", "clone", "--quiet", GIT_REPO_URL, REPO_DIR], check=True)
os.makedirs(f"{REPO_DIR}/daily_stories", exist_ok=True)
import shutil
shutil.copy(filename, f"{REPO_DIR}/daily_stories/{filename}")
subprocess.run(["git", "-C", REPO_DIR, "add", "daily_stories/"], check=True)
subprocess.run(["git", "-C", REPO_DIR, "commit", "-m", "Test daily story update"], check=True)
subprocess.run(["git", "-C", REPO_DIR, "push", "--quiet"], check=True)
print("✅ Test upload complete")
