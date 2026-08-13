#!/usr/bin/env python3
"""
Read story URLs from daily_stories/stories_YYYYMMDD.txt (clone locally),
fetch each story page, extract div.panel-body content,
write to a Word doc, and push to GitHub.
"""

import os
import sys
import subprocess
import glob

# Ensure dependencies are installed
subprocess.run([sys.executable, "-m", "pip", "install", "--quiet", "beautifulsoup4", "python-docx", "requests"],
               stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)

import requests
from bs4 import BeautifulSoup
from docx import Document
from datetime import datetime
import shutil

BASE_URL = "https://shahvani.com"
REPO_DIR = "/tmp/hermes-backup-git"
GITHUB_USER = "chertopert1981"
GITHUB_REPO = "chertopert1981/hermes-backup"
GITHUB_PAT = "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"
GIT_REPO_URL = f"https://{GITHUB_USER}:{GITHUB_PAT}@github.com/{GITHUB_REPO}.git"

def main():
    date_str = datetime.now().strftime("%Y%m%d")

    # Clean local changes before pull
    if os.path.exists(REPO_DIR):
        for path in glob.glob(f"{REPO_DIR}/daily_stories/stories_*"):
            try:
                os.remove(path)
            except Exception:
                pass
        subprocess.run(["git", "-C", REPO_DIR, "checkout", "--", "daily_stories/"], check=False)
        subprocess.run(["git", "-C", REPO_DIR, "clean", "-fd", "daily_stories/"], check=False)
        subprocess.run(["git", "-C", REPO_DIR, "pull", "--rebase"], check=True)
    else:
        subprocess.run(["git", "clone", "--quiet", GIT_REPO_URL, REPO_DIR], check=True)
        subprocess.run(["git", "-C", REPO_DIR, "checkout", "main"], check=True)

    txt_path = f"{REPO_DIR}/daily_stories/stories_{date_str}.txt"
    if not os.path.exists(txt_path):
        print(f"\u274c File not found: {txt_path}")
        return 1

    print(f"\U0001f4e5 Reading story list from local clone: {txt_path}")
    with open(txt_path, "r") as f:
        links = [l.strip() for l in f if l.strip()]

    print(f"\U0001f4da Found {len(links)} story links")

    if not links:
        print("\u26a0\ufe0f No links found in the file.")
        return 1

    # Create Word document
    doc = Document()
    doc.add_heading(f"Shahvani Stories - {date_str}", 0)

    count = 0
    for i, link in enumerate(links, 1):
        print(f"  [{i}/{len(links)}] Fetching: {link[:60]}...")
        try:
            page_resp = requests.get(link, headers={"User-Agent": "Mozilla/5.0"}, timeout=30)
            page_soup = BeautifulSoup(page_resp.text, "html.parser")

            panel_body = page_soup.select_one("div.panel-body")
            if panel_body:
                # Remove everything after div#loginorregister
                login_div = panel_body.select_one("div#loginorregister")
                if login_div:
                    for sibling in list(login_div.next_siblings):
                        sibling.extract()
                    login_div.extract()
                content = panel_body.get_text(separator="\n", strip=True)
            else:
                content = "[mohavate yafat nashod]"

            doc.add_paragraph(f"\U0001f517 {link}")
            doc.add_paragraph(content)
            doc.add_paragraph("\u2500" * 50)
            count += 1

        except Exception as e:
            print(f"    \u274c Error: {e}")
            doc.add_paragraph(f"\U0001f517 {link}")
            doc.add_paragraph(f"[kharabar dar daryaft: {e}]")
            doc.add_paragraph("\u2500" * 50)

    # Save docx
    filename = f"stories_{date_str}.docx"
    doc.save(filename)
    print(f"\u2705 Created {filename} with {count} stories")

    # Copy to repo and push
    os.makedirs(f"{REPO_DIR}/daily_stories", exist_ok=True)
    shutil.copy(filename, f"{REPO_DIR}/daily_stories/{filename}")

    subprocess.run(["git", "-C", REPO_DIR, "add", f"daily_stories/{filename}"], check=True)
    subprocess.run(["git", "-C", REPO_DIR, "commit", "-m", f"Stories docx {date_str}"], check=True)
    subprocess.run(["git", "-C", REPO_DIR, "push", "--quiet"], check=True)
    print("\u2705 Pushed to GitHub daily_stories/")

    # Cleanup
    os.remove(filename)

if __name__ == "__main__":
    main()
